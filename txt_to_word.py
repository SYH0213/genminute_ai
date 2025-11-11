#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
텍스트 파일을 Word 파일로 변환 (Q&A 색상 구분)
질문(Q)은 파란색, 답변(A)은 검정색으로 표시
"""

import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_word_from_qa(input_file, output_file=None):
    """Q&A 텍스트 파일을 Word 문서로 변환"""

    if output_file is None:
        output_file = input_file.replace('.txt', '.docx')

    print(f"📄 파일 읽기: {input_file}")

    try:
        # 텍스트 파일 읽기
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Word 문서 생성
        doc = Document()

        # 제목 추가
        title = doc.add_heading('심사위원 예상 질문 및 답변', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 부제목
        subtitle = doc.add_paragraph('Minute AI 프로젝트 발표용')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(12)
        subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # 빈 줄

        # Q&A 패턴 찾기
        # Q1. 질문?
        # A1. 답변.
        lines = content.split('\n')

        for line in lines:
            line = line.strip()

            # 구분선 건너뛰기
            if line.startswith('====') or line.startswith('----'):
                continue

            # 제목 건너뛰기 (이미 추가함)
            if '심사위원 예상 질문' in line or 'Minute AI' in line:
                continue

            # 빈 줄
            if not line:
                continue

            # Q로 시작하는 줄 (질문)
            if re.match(r'^Q\d+\.', line):
                # 질문 번호와 내용 분리
                match = re.match(r'^(Q\d+)\.\s*(.+)$', line)
                if match:
                    q_num = match.group(1)
                    q_text = match.group(2)

                    # 단락 추가
                    p = doc.add_paragraph()

                    # 질문 번호 (굵게, 파란색)
                    run_num = p.add_run(f'{q_num}. ')
                    run_num.font.bold = True
                    run_num.font.size = Pt(11)
                    run_num.font.color.rgb = RGBColor(0, 51, 153)  # 진한 파란색

                    # 질문 내용 (파란색)
                    run_text = p.add_run(q_text)
                    run_text.font.size = Pt(11)
                    run_text.font.color.rgb = RGBColor(0, 102, 204)  # 파란색

            # A로 시작하는 줄 (답변)
            elif re.match(r'^A\d+\.', line):
                # 답변 번호와 내용 분리
                match = re.match(r'^(A\d+)\.\s*(.+)$', line)
                if match:
                    a_num = match.group(1)
                    a_text = match.group(2)

                    # 단락 추가
                    p = doc.add_paragraph()

                    # 답변 번호 (굵게, 녹색)
                    run_num = p.add_run(f'{a_num}. ')
                    run_num.font.bold = True
                    run_num.font.size = Pt(11)
                    run_num.font.color.rgb = RGBColor(0, 102, 0)  # 진한 녹색

                    # 답변 내용 (검정색)
                    run_text = p.add_run(a_text)
                    run_text.font.size = Pt(11)
                    run_text.font.color.rgb = RGBColor(0, 0, 0)  # 검정색

                    # 답변 뒤 여백
                    doc.add_paragraph()

        # Word 파일 저장
        doc.save(output_file)

        print(f"✅ Word 파일 생성 완료: {output_file}")
        print(f"📊 질문은 파란색, 답변은 검정색으로 표시되었습니다.")

        return True

    except ImportError:
        print("❌ python-docx 라이브러리가 설치되어 있지 않습니다.")
        print("\n설치 방법:")
        print("  pip install python-docx")
        return False

    except Exception as e:
        print(f"❌ 변환 실패: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    if len(sys.argv) < 2:
        print("사용법: python txt_to_word.py <input.txt> [output.docx]")
        print("\n예시:")
        print("  python txt_to_word.py 심사위원_예상질문_50선_한줄답변.txt")
        print("  python txt_to_word.py input.txt output.docx")
        return

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    create_word_from_qa(input_file, output_file)


if __name__ == "__main__":
    main()
